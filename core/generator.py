"""Main document generator for the DocStyle Transformer.

Takes a :class:`DocumentTree` intermediate representation and a
:class:`DesignSystem` configuration, then produces a fully styled ``.docx``
file.  Every element type in the IR is rendered into python-docx content with
typography, colour, and layout derived from the design system.

Usage::

    from core.generator import DocumentGenerator

    gen = DocumentGenerator()
    output = gen.generate(tree, "output/report.docx")
"""

from __future__ import annotations

import io
import logging
from pathlib import Path
from typing import Any, TYPE_CHECKING

from docx import Document as new_docx
from docx.shared import Pt, Twips, RGBColor, Inches, Emu, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from lxml import etree

from core.mapper import DesignSystem, StyleMapper
from core.models import (
    Callout,
    CalloutType,
    ContentElement,
    DocumentMetadata,
    DocumentTree,
    Image,
    ListBlock,
    ListItem,
    ListType,
    PageBreak,
    Paragraph,
    Section,
    Step,
    StepsBlock,
    Table,
    TextRun,
)
from core.cover import CoverGenerator
from core.toc import TocGenerator

if TYPE_CHECKING:
    from docx.document import Document
    from docx.text.paragraph import Paragraph as DocxParagraph
    from docx.text.run import Run

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _hex_to_rgb(hex_color: str) -> RGBColor:
    """Convert a ``#RRGGBB`` hex string to a python-docx *RGBColor*.

    The leading ``#`` is optional.
    """
    hex_color = hex_color.lstrip("#")
    if len(hex_color) != 6:
        raise ValueError(f"Invalid hex color: #{hex_color}")
    r = int(hex_color[0:2], 16)
    g = int(hex_color[2:4], 16)
    b = int(hex_color[4:6], 16)
    return RGBColor(r, g, b)


def _set_cell_shading(cell: Any, hex_color: str) -> None:
    """Apply a solid background fill to a table *cell*.

    Parameters
    ----------
    cell:
        A python-docx table cell.
    hex_color:
        A ``#RRGGBB`` hex colour string.
    """
    color_val = hex_color.lstrip("#")
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), color_val)
    shading_elm.set(qn("w:val"), "clear")
    shading_elm.set(qn("w:color"), "auto")
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _set_cell_padding(cell: Any, top: int = 0, bottom: int = 0,
                      left: int = 0, right: int = 0) -> None:
    """Set individual cell margins (padding) in DXA.

    Parameters
    ----------
    cell:
        A python-docx table cell.
    top, bottom, left, right:
        Padding values in DXA (twentieths of a point).
    """
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for side, value in [("top", top), ("bottom", bottom),
                        ("start", left), ("end", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(value))
        el.set(qn("w:type"), "dxa")
        tc_mar.append(el)
    tc_pr.append(tc_mar)


def _set_cell_borders(cell: Any, top: str | None = None,
                      bottom: str | None = None, left: str | None = None,
                      right: str | None = None, size: int = 4,
                      color: str = "D2D2D7") -> None:
    """Set borders on an individual table *cell*.

    Each side parameter, if provided, should be a border style string
    (e.g. ``"single"``).  *size* is in eighth-points, *color* is a bare
    hex string without ``#``.
    """
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom),
                      ("start", left), ("end", right)]:
        if val is not None:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), val)
            el.set(qn("w:sz"), str(size))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color.lstrip("#"))
            tc_borders.append(el)
    tc_pr.append(tc_borders)


def _set_table_borders(table: Any, color: str = "D2D2D7",
                       size: int = 4) -> None:
    """Apply uniform thin borders to all edges of a docx *table*.

    Parameters
    ----------
    table:
        A python-docx Table object.
    color:
        Bare hex colour string (no ``#``).
    size:
        Border width in eighth-points.
    """
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color.lstrip("#"))
        borders.append(el)
    tbl_pr.append(borders)


def _add_bottom_border(para: DocxParagraph, color: str,
                       size_pt: int = 6) -> None:
    """Add a bottom border to a paragraph (accent bar effect).

    Parameters
    ----------
    para:
        A python-docx Paragraph.
    color:
        Hex colour string (with or without ``#``).
    size_pt:
        Border thickness in points.
    """
    p_pr = para._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size_pt * 8))  # 1/8 pt units
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color.lstrip("#"))
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _add_top_border(para: DocxParagraph, color: str,
                    size_pt: int = 1) -> None:
    """Add a top border to a paragraph.

    Parameters
    ----------
    para:
        A python-docx Paragraph.
    color:
        Hex colour string (with or without ``#``).
    size_pt:
        Border thickness in points.
    """
    p_pr = para._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), str(size_pt * 8))
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), color.lstrip("#"))
    p_bdr.append(top)
    p_pr.append(p_bdr)


# ---------------------------------------------------------------------------
# DocumentGenerator
# ---------------------------------------------------------------------------


class DocumentGenerator:
    """Renders a :class:`DocumentTree` into a styled ``.docx`` file.

    The generator orchestrates cover-page creation, table-of-contents
    generation, section rendering, and header/footer setup.  All visual
    parameters are drawn from the :class:`DesignSystem`.

    Parameters
    ----------
    design_system : DesignSystem or None, optional
        An already-loaded design system instance.  When ``None`` a default
        :class:`DesignSystem` is created (optionally influenced by
        *theme_path*).
    theme_path : str or None, optional
        Path to a theme overlay YAML applied on top of the base design-system
        configuration.  Only used when *design_system* is ``None``.
    """

    def __init__(
        self,
        design_system: DesignSystem | None = None,
        theme_path: str | None = None,
    ) -> None:
        if design_system is not None:
            self._design = design_system
        elif theme_path is not None:
            self._design = DesignSystem(theme_path=theme_path)
        else:
            self._design = DesignSystem()

        self._mapper = StyleMapper(self._design)
        self._cover_gen = CoverGenerator(self._design)
        self._toc_gen = TocGenerator(self._design)

        logger.info("DocumentGenerator initialised")

    # ── Public API ────────────────────────────────────────────────────

    def generate(
        self,
        tree: DocumentTree,
        output_path: str,
        options: dict | None = None,
    ) -> str:
        """Generate a ``.docx`` file from the document *tree*.

        Parameters
        ----------
        tree : DocumentTree
            The intermediate representation of the parsed document.
        output_path : str
            Destination file path for the generated ``.docx``.
        options : dict or None, optional
            Generation options.  Supported keys:

            - ``generate_cover`` (bool, default ``True``): include a cover page.
            - ``generate_toc`` (bool, default ``True``): include a table of
              contents.
            - ``number_sections`` (bool, default ``True``): prefix H1 headings
              with "Section XX".
            - ``header_footer`` (bool, default ``True``): add running
              header and footer.
            - ``cover_title_override`` (str or None): replace the metadata
              title on the cover page.

        Returns
        -------
        str
            The absolute path to the generated file.
        """
        opts = options or {}
        generate_cover: bool = opts.get("generate_cover", True)
        generate_toc: bool = opts.get("generate_toc", True)
        number_sections: bool = opts.get("number_sections", True)
        header_footer: bool = opts.get("header_footer", True)
        cover_title_override: str | None = opts.get("cover_title_override")

        logger.info(
            "Generating document: cover=%s, toc=%s, numbered=%s, hf=%s",
            generate_cover, generate_toc, number_sections, header_footer,
        )

        # -- 1. Create the document and configure page layout ---------------
        doc = new_docx()
        self._setup_page(doc)

        metadata = tree.metadata

        # -- 2. Cover page --------------------------------------------------
        if generate_cover:
            cover_meta = metadata
            if cover_title_override:
                cover_meta = DocumentMetadata(
                    title=cover_title_override,
                    author=metadata.author,
                    date=metadata.date,
                    version=metadata.version,
                    reference=metadata.reference,
                )
            self._cover_gen.generate(doc, cover_meta)
            logger.debug("Cover page appended")

        # -- 3. Table of contents -------------------------------------------
        if generate_toc:
            self._toc_gen.generate(doc, tree)
            logger.debug("Table of contents appended")

        # -- 4. Preamble (content before first section) ---------------------
        if tree.preamble:
            logger.debug("Rendering %d preamble element(s)", len(tree.preamble))
            for element in tree.preamble:
                self._render_element(doc, element)

        # -- 5. Sections ----------------------------------------------------
        for section in tree.sections:
            heading_style = self._mapper.map_heading(section.level)
            self._render_heading(doc, section, heading_style,
                                 number_sections=number_sections)

            for element in section.children:
                self._render_element(doc, element)

        # -- 6. Headers and footers -----------------------------------------
        if header_footer:
            self._setup_header_footer(doc, metadata)
            logger.debug("Headers and footers configured")

        # -- 7. Save --------------------------------------------------------
        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out))

        resolved_path = str(out.resolve())
        logger.info("Document saved to %s", resolved_path)
        return resolved_path

    # ── Page setup ────────────────────────────────────────────────────

    def _setup_page(self, doc: Document) -> None:
        """Configure page dimensions and margins from the design system."""
        page_cfg = self._design.get_page_config()
        section = doc.sections[0]

        section.page_width = Twips(page_cfg.get("width", 12240))
        section.page_height = Twips(page_cfg.get("height", 15840))

        margins = page_cfg.get("margins", {})
        section.top_margin = Twips(margins.get("top", 1440))
        section.right_margin = Twips(margins.get("right", 1440))
        section.bottom_margin = Twips(margins.get("bottom", 1440))
        section.left_margin = Twips(margins.get("left", 1440))

        logger.debug(
            "Page: %dx%d DXA, margins T=%d R=%d B=%d L=%d",
            page_cfg.get("width", 12240),
            page_cfg.get("height", 15840),
            margins.get("top", 1440),
            margins.get("right", 1440),
            margins.get("bottom", 1440),
            margins.get("left", 1440),
        )

    # ── Element dispatch ──────────────────────────────────────────────

    def _render_element(self, doc: Document, element: ContentElement) -> None:
        """Route a content element to its type-specific renderer."""
        style = self._mapper.map_element(element)

        match element:
            case Paragraph():
                self._render_paragraph(doc, element, style)
            case Table():
                self._render_table(doc, element, style)
            case Image():
                self._render_image(doc, element, style)
            case Callout():
                self._render_callout(doc, element, style)
            case ListBlock():
                self._render_list(doc, element, style)
            case StepsBlock():
                self._render_steps(doc, element, style)
            case PageBreak():
                self._render_page_break(doc)
            case _:
                logger.warning(
                    "Unsupported element type: %s", type(element).__name__
                )

    # ── Heading rendering ─────────────────────────────────────────────

    def _render_heading(
        self,
        doc: Document,
        section: Section,
        style: dict[str, Any],
        *,
        number_sections: bool = True,
    ) -> None:
        """Render a section heading into the document.

        For **level 1** headings the output is:

        1. A "Section XX" label paragraph in accent_blue.
        2. The heading title in bold black.
        3. An accent bar (bottom-border paragraph).

        For **level 2 / level 3** headings a simpler single paragraph with
        the appropriate size and colour is produced.

        Parameters
        ----------
        doc:
            The python-docx Document.
        section:
            The :class:`Section` whose heading is being rendered.
        style:
            Style dictionary from :meth:`StyleMapper.map_heading`.
        number_sections:
            Whether to prefix level-1 headings with "Section XX".
        """
        level = section.level
        font_name = style.get("font", self._design.get_font("display"))
        size_hp = style.get("size", 40)
        size_pt = size_hp / 2.0
        color_hex = self._design.resolve_color(style.get("color", "black"))
        spacing_before = style.get("spacing_before", 480)
        spacing_after = style.get("spacing_after", 200)
        keep_with_next = style.get("keep_with_next", True)

        if level == 1:
            self._render_heading_level1(
                doc, section, font_name=font_name, size_pt=size_pt,
                spacing_before=spacing_before, spacing_after=spacing_after,
                keep_with_next=keep_with_next,
                number_sections=number_sections,
            )
        else:
            self._render_heading_sub(
                doc, section, font_name=font_name, size_pt=size_pt,
                color_hex=color_hex, spacing_before=spacing_before,
                spacing_after=spacing_after, keep_with_next=keep_with_next,
            )

    def _render_heading_level1(
        self,
        doc: Document,
        section: Section,
        *,
        font_name: str,
        size_pt: float,
        spacing_before: int,
        spacing_after: int,
        keep_with_next: bool,
        number_sections: bool,
    ) -> None:
        """Render a level-1 section heading with accent label and bar."""
        accent_hex = self._design.resolve_color("accent_blue")
        black_hex = self._design.resolve_color("black")
        section_number = section.number if section.number is not None else 0

        # -- "Section XX" label --
        if number_sections:
            label_para = doc.add_paragraph()
            label_para.paragraph_format.left_indent = Twips(0)
            label_para.paragraph_format.first_line_indent = Twips(0)
            label_para.paragraph_format.space_before = Twips(spacing_before)
            label_para.paragraph_format.space_after = Twips(0)
            label_para.paragraph_format.keep_with_next = True

            label_run = label_para.add_run(f"Section {section_number:02d}")
            label_run.bold = True
            label_run.font.size = Pt(size_pt * 0.6)
            label_run.font.color.rgb = _hex_to_rgb(accent_hex)
            label_run.font.name = font_name

        # -- Title --
        title_para = doc.add_paragraph()
        title_para.paragraph_format.left_indent = Twips(0)
        title_para.paragraph_format.first_line_indent = Twips(0)
        if not number_sections:
            title_para.paragraph_format.space_before = Twips(spacing_before)
        else:
            title_para.paragraph_format.space_before = Twips(80)
        title_para.paragraph_format.space_after = Twips(0)
        title_para.paragraph_format.keep_with_next = True

        title_run = title_para.add_run(section.heading)
        title_run.bold = True
        title_run.font.size = Pt(size_pt)
        title_run.font.color.rgb = _hex_to_rgb(black_hex)
        title_run.font.name = font_name

        # -- Accent bar --
        bar_para = doc.add_paragraph()
        bar_para.paragraph_format.left_indent = Twips(0)
        bar_para.paragraph_format.first_line_indent = Twips(0)
        bar_para.paragraph_format.space_before = Twips(80)
        bar_para.paragraph_format.space_after = Twips(spacing_after)

        accent_bar_style = self._design.get_component_style("accent_bar")
        bar_height = accent_bar_style.get("border_bottom_size", 6)
        _add_bottom_border(bar_para, accent_hex, size_pt=bar_height)

        # Constrain bar width via right indent.
        usable = self._design.usable_width
        bar_width = 2400  # DXA - standard accent bar width
        right_indent = max(usable - bar_width, 0)
        bar_para.paragraph_format.right_indent = Twips(right_indent)

        logger.debug(
            "Rendered H1: '%s' (section %02d)", section.heading, section_number
        )

    def _render_heading_sub(
        self,
        doc: Document,
        section: Section,
        *,
        font_name: str,
        size_pt: float,
        color_hex: str,
        spacing_before: int,
        spacing_after: int,
        keep_with_next: bool,
    ) -> None:
        """Render a level-2 or level-3 heading as a simple styled paragraph."""
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Twips(0)
        para.paragraph_format.first_line_indent = Twips(0)
        para.paragraph_format.space_before = Twips(spacing_before)
        para.paragraph_format.space_after = Twips(spacing_after)
        para.paragraph_format.keep_with_next = keep_with_next

        run = para.add_run(section.heading)
        run.bold = True
        run.font.size = Pt(size_pt)
        run.font.color.rgb = _hex_to_rgb(color_hex)
        run.font.name = font_name

        logger.debug(
            "Rendered H%d: '%s'", section.level, section.heading
        )

    # ── Paragraph rendering ───────────────────────────────────────────

    def _render_paragraph(
        self,
        doc: Document,
        paragraph: Paragraph,
        style: dict[str, Any],
    ) -> None:
        """Render a body paragraph with styled text runs.

        Each :class:`TextRun` in the paragraph is converted to a python-docx
        run with font, size, colour, and inline-formatting attributes applied.
        """
        body_font = style.get("font", self._design.get_font("body"))
        body_size_hp = style.get("size", 21)
        body_color_hex = self._design.resolve_color(
            style.get("color", "dark_gray")
        )
        line_spacing_twips = style.get("line_spacing")
        spacing_after_twips = style.get("spacing_after", 160)

        docx_para = doc.add_paragraph()

        # Paragraph-level formatting — explicitly zero out indentation so that
        # the user's Normal.dotm template cannot inject unexpected indents.
        docx_para.paragraph_format.left_indent = Twips(0)
        docx_para.paragraph_format.first_line_indent = Twips(0)
        if line_spacing_twips:
            docx_para.paragraph_format.line_spacing = Twips(line_spacing_twips)
        docx_para.paragraph_format.space_after = Twips(spacing_after_twips)

        if not paragraph.runs:
            # Empty paragraph -- just apply base style to an empty run.
            empty_run = docx_para.add_run("")
            empty_run.font.name = body_font
            empty_run.font.size = Pt(body_size_hp / 2.0)
            empty_run.font.color.rgb = _hex_to_rgb(body_color_hex)
            return

        for text_run in paragraph.runs:
            run = docx_para.add_run(text_run.text)

            # Font name -- use run-specific override or body default.
            run.font.name = body_font

            # Font size -- always use the design-system body size to ensure
            # consistent typography.  The original run size (if any) is
            # intentionally ignored because the goal of the transformer is
            # to apply a uniform design system.
            run.font.size = Pt(body_size_hp / 2.0)

            # Colour -- prefer run-specific colour, fall back to body colour.
            if text_run.color:
                run.font.color.rgb = _hex_to_rgb(
                    self._design.resolve_color(text_run.color)
                )
            else:
                run.font.color.rgb = _hex_to_rgb(body_color_hex)

            # Inline formatting
            if text_run.bold:
                run.bold = True
            if text_run.italic:
                run.italic = True
            if text_run.underline:
                run.underline = True
            if text_run.strikethrough:
                run.font.strike = True

            # Hyperlinks are stored on the TextRun but need low-level XML
            # manipulation to render as clickable links in OOXML.  We add a
            # visual underline + blue colour as a hint; a full hyperlink
            # relationship is complex and deferred to a future enhancement.
            if text_run.hyperlink:
                run.underline = True
                run.font.color.rgb = _hex_to_rgb(
                    self._design.resolve_color("accent_blue")
                )

        logger.debug(
            "Rendered paragraph (%d run(s)): '%s'",
            len(paragraph.runs),
            paragraph.text[:60] if paragraph.text else "",
        )

    # ── Table rendering ───────────────────────────────────────────────

    def _render_table(
        self,
        doc: Document,
        table: Table,
        style: dict[str, Any],
    ) -> None:
        """Render a table with styled header row, zebra striping, and borders.

        The header row receives an accent_blue background with white text.
        Data rows alternate between white and a very light grey background.
        """
        header_bg_hex = self._design.resolve_color(
            style.get("header_bg", "accent_blue")
        )
        header_text_hex = self._design.resolve_color(
            style.get("header_text_color", "white")
        )
        header_size_hp = style.get("header_font_size", 21)
        header_bold = style.get("header_bold", True)

        body_size_hp = style.get("body_font_size", 21)
        body_color_hex = self._design.resolve_color(
            style.get("body_color", "dark_gray")
        )
        body_font = style.get("font", self._design.get_font("body"))

        zebra_stripe = style.get("zebra_stripe", True)
        zebra_hex = self._design.resolve_color(
            style.get("zebra_color", "bg_light")
        )
        border_color_hex = self._design.resolve_color(
            style.get("row_border_color", "very_light_gray")
        )
        border_size = style.get("row_border_size", 1)

        cell_pad_top = style.get("cell_padding_top", 80)
        cell_pad_bottom = style.get("cell_padding_bottom", 80)
        cell_pad_left = style.get("cell_padding_left", 120)
        cell_pad_right = style.get("cell_padding_right", 120)

        # Determine dimensions
        col_count = len(table.headers) if table.headers else (
            len(table.rows[0]) if table.rows else 0
        )
        if col_count == 0:
            logger.warning("Skipping empty table (no columns)")
            return

        total_rows = (1 if table.headers else 0) + len(table.rows)
        docx_table = doc.add_table(rows=total_rows, cols=col_count)
        docx_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Apply uniform thin borders
        _set_table_borders(
            docx_table,
            color=border_color_hex.lstrip("#"),
            size=border_size * 8,
        )

        current_row = 0

        # -- Header row ----------------------------------------------------
        if table.headers:
            for col_idx, header_text in enumerate(table.headers):
                cell = docx_table.cell(current_row, col_idx)
                _set_cell_shading(cell, header_bg_hex)
                _set_cell_padding(cell, top=cell_pad_top, bottom=cell_pad_bottom,
                                  left=cell_pad_left, right=cell_pad_right)

                cell_para = cell.paragraphs[0]
                cell_para.paragraph_format.space_before = Twips(0)
                cell_para.paragraph_format.space_after = Twips(0)

                # Use header_runs if available, otherwise plain text.
                if table.header_runs and col_idx < len(table.header_runs):
                    for text_run in table.header_runs[col_idx]:
                        run = cell_para.add_run(text_run.text)
                        run.bold = header_bold
                        run.font.size = Pt(header_size_hp / 2.0)
                        run.font.color.rgb = _hex_to_rgb(header_text_hex)
                        run.font.name = body_font
                else:
                    run = cell_para.add_run(header_text)
                    run.bold = header_bold
                    run.font.size = Pt(header_size_hp / 2.0)
                    run.font.color.rgb = _hex_to_rgb(header_text_hex)
                    run.font.name = body_font

            current_row += 1

        # -- Data rows -----------------------------------------------------
        for row_idx, row_data in enumerate(table.rows):
            is_even_row = row_idx % 2 == 0

            for col_idx in range(col_count):
                cell_text = row_data[col_idx] if col_idx < len(row_data) else ""
                cell = docx_table.cell(current_row, col_idx)

                # Zebra striping (even rows get the light background)
                if zebra_stripe and is_even_row:
                    _set_cell_shading(cell, zebra_hex)

                _set_cell_padding(cell, top=cell_pad_top, bottom=cell_pad_bottom,
                                  left=cell_pad_left, right=cell_pad_right)

                cell_para = cell.paragraphs[0]
                cell_para.paragraph_format.space_before = Twips(0)
                cell_para.paragraph_format.space_after = Twips(0)

                # Use cell_runs if available, otherwise plain text.
                if (table.cell_runs
                        and row_idx < len(table.cell_runs)
                        and col_idx < len(table.cell_runs[row_idx])):
                    for text_run in table.cell_runs[row_idx][col_idx]:
                        run = cell_para.add_run(text_run.text)
                        run.font.size = Pt(body_size_hp / 2.0)
                        run.font.color.rgb = _hex_to_rgb(body_color_hex)
                        run.font.name = body_font
                        if text_run.bold:
                            run.bold = True
                        if text_run.italic:
                            run.italic = True
                else:
                    run = cell_para.add_run(cell_text)
                    run.font.size = Pt(body_size_hp / 2.0)
                    run.font.color.rgb = _hex_to_rgb(body_color_hex)
                    run.font.name = body_font

            current_row += 1

        logger.debug(
            "Rendered table: %d col(s), %d data row(s), headers=%s",
            col_count, len(table.rows), bool(table.headers),
        )

    # ── Image rendering ───────────────────────────────────────────────

    def _render_image(
        self,
        doc: Document,
        image: Image,
        style: dict[str, Any],
    ) -> None:
        """Render an embedded image, scaling to fit within the usable width.

        The image binary data is read from :attr:`Image.data` via an
        in-memory buffer.  If the image's native width exceeds the page's
        usable width it is proportionally scaled down.
        """
        if not image.data:
            logger.warning("Skipping image with no data: %s", image.filename)
            return

        max_width_dxa = style.get("max_width", self._design.usable_width)
        # Convert DXA to EMU (1 DXA = 635 EMU)
        max_width_emu = max_width_dxa * 635

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Twips(120)
        para.paragraph_format.space_after = Twips(120)

        stream = io.BytesIO(image.data)
        run = para.add_run()

        # Determine desired width.  If the image has explicit dimensions
        # (in DXA) use those; otherwise let python-docx determine the
        # native size and we will post-check against the maximum.
        if image.width and image.width > 0:
            img_width_emu = image.width * 635
            if image.height and image.height > 0:
                img_height_emu = image.height * 635
            else:
                img_height_emu = None

            # Scale down if wider than usable area
            if img_width_emu > max_width_emu:
                scale = max_width_emu / img_width_emu
                img_width_emu = int(max_width_emu)
                if img_height_emu is not None:
                    img_height_emu = int(img_height_emu * scale)

            if img_height_emu is not None:
                run.add_picture(stream, width=Emu(img_width_emu),
                                height=Emu(img_height_emu))
            else:
                run.add_picture(stream, width=Emu(img_width_emu))
        else:
            # No explicit dimensions -- insert at native size, then check.
            inline_shape = run.add_picture(stream)
            if inline_shape.width > max_width_emu:
                scale = max_width_emu / inline_shape.width
                inline_shape.width = int(max_width_emu)
                inline_shape.height = int(inline_shape.height * scale)

        # Alt text / caption
        if image.alt_text:
            caption_style = style.get("caption_style", {})
            cap_para = doc.add_paragraph()
            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_para.paragraph_format.space_before = Twips(40)
            cap_para.paragraph_format.space_after = Twips(
                caption_style.get("spacing_after", 80)
            )
            cap_run = cap_para.add_run(image.alt_text)
            cap_run.italic = True
            cap_run.font.size = Pt(
                caption_style.get("size", 18) / 2.0
            )
            cap_color = caption_style.get("color")
            if cap_color:
                cap_run.font.color.rgb = _hex_to_rgb(
                    self._design.resolve_color(cap_color)
                )
            cap_font = caption_style.get("font", self._design.get_font("body"))
            cap_run.font.name = cap_font

        logger.debug("Rendered image: %s", image.filename)

    # ── Callout rendering ─────────────────────────────────────────────

    def _render_callout(
        self,
        doc: Document,
        callout: Callout,
        style: dict[str, Any],
    ) -> None:
        """Render a callout box as a single-cell table with a coloured left border.

        Info/note/tip callouts use accent_blue styling; warning callouts use
        the warning_border colour and warning_bg background.
        """
        border_color_hex = self._design.resolve_color(
            style.get("border_left_color", "accent_blue")
        )
        bg_hex = self._design.resolve_color(
            style.get("background", "bg_light")
        )
        border_size = style.get("border_left_size", 12)

        title_size_hp = style.get("title_size", 21)
        title_bold = style.get("title_bold", True)
        title_color_hex = self._design.resolve_color(
            style.get("title_color", "accent_blue")
        )

        body_size_hp = style.get("body_size", 21)
        body_color_hex = self._design.resolve_color(
            style.get("body_color", "dark_gray")
        )
        body_font = style.get("font", self._design.get_font("body"))

        pad_top = style.get("padding_top", 120)
        pad_bottom = style.get("padding_bottom", 120)
        pad_left = style.get("padding_left", 240)
        pad_right = style.get("padding_right", 240)

        spacing_before = style.get("spacing_before", 200)
        spacing_after = style.get("spacing_after", 200)

        # Create a single-row, single-column table to simulate the callout box
        callout_table = doc.add_table(rows=1, cols=1)
        callout_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        cell = callout_table.cell(0, 0)
        _set_cell_shading(cell, bg_hex)
        _set_cell_padding(cell, top=pad_top, bottom=pad_bottom,
                          left=pad_left, right=pad_right)

        # Apply a thick left border and remove other borders
        _set_cell_borders(
            cell,
            left="single",
            top="none",
            bottom="none",
            right="none",
            size=border_size * 8,  # convert pt to eighth-points
            color=border_color_hex,
        )

        # Clear the default paragraph
        cell_para = cell.paragraphs[0]
        cell_para.paragraph_format.space_before = Twips(0)
        cell_para.paragraph_format.space_after = Twips(80)

        # -- Title run --
        title_text = callout.title
        if not title_text:
            # Default titles based on callout type
            type_titles = {
                CalloutType.INFO: "Information",
                CalloutType.WARNING: "Attention",
                CalloutType.NOTE: "Note",
                CalloutType.TIP: "Conseil",
            }
            title_text = type_titles.get(callout.callout_type, "Note")

        title_run = cell_para.add_run(title_text)
        title_run.bold = title_bold
        title_run.font.size = Pt(title_size_hp / 2.0)
        title_run.font.color.rgb = _hex_to_rgb(title_color_hex)
        title_run.font.name = body_font

        # -- Body text --
        if callout.body or callout.body_runs:
            body_para = cell.add_paragraph()
            body_para.paragraph_format.space_before = Twips(40)
            body_para.paragraph_format.space_after = Twips(0)

            if callout.body_runs:
                for text_run in callout.body_runs:
                    run = body_para.add_run(text_run.text)
                    run.font.size = Pt(body_size_hp / 2.0)
                    run.font.name = body_font
                    if text_run.bold:
                        run.bold = True
                    if text_run.italic:
                        run.italic = True
                    if text_run.color:
                        run.font.color.rgb = _hex_to_rgb(
                            self._design.resolve_color(text_run.color)
                        )
                    else:
                        run.font.color.rgb = _hex_to_rgb(body_color_hex)
            else:
                run = body_para.add_run(callout.body)
                run.font.size = Pt(body_size_hp / 2.0)
                run.font.color.rgb = _hex_to_rgb(body_color_hex)
                run.font.name = body_font

        # Add spacing around the callout table via a wrapper paragraph trick:
        # we set the spacing on the paragraph *before* the table was added.
        # Since python-docx doesn't directly support table-level spacing,
        # we manipulate the underlying XML tblPr element.
        tbl_pr = callout_table._tbl.tblPr
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            callout_table._tbl.insert(0, tbl_pr)

        # Remove default table borders (the cell-level left border is enough)
        tbl_borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "none")
            el.set(qn("w:sz"), "0")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "auto")
            tbl_borders.append(el)
        tbl_pr.append(tbl_borders)

        # Spacing before/after via tblpPr (not universally supported but
        # helps in Word).  We also add an empty spacer paragraph before.
        spacer_before = doc.paragraphs[-1] if doc.paragraphs else None
        # Move the table after its natural position -- spacing is handled
        # by adding small paragraphs.
        space_para = doc.add_paragraph()
        space_para.paragraph_format.space_before = Twips(0)
        space_para.paragraph_format.space_after = Twips(spacing_after)
        # Make it visually empty with a zero-size font.
        tiny_run = space_para.add_run("")
        tiny_run.font.size = Pt(1)

        logger.debug(
            "Rendered %s callout: '%s'",
            callout.callout_type.value, title_text,
        )

    # ── List rendering ────────────────────────────────────────────────

    def _render_list(
        self,
        doc: Document,
        list_block: ListBlock,
        style: dict[str, Any],
    ) -> None:
        """Render a bullet or numbered list.

        Bullets are rendered with a ``\\u2022`` (bullet) prefix; numbered
        items use ``1.``, ``2.`` etc.  Nesting is achieved through left
        indentation.
        """
        is_numbered = list_block.list_type == ListType.NUMBERED
        indent_per_level = style.get("indent_per_level", 360)
        bullet_color_hex = self._design.resolve_color(
            style.get("bullet_color",
                       style.get("number_color", "accent_blue"))
        )
        text_size_hp = style.get("text_size", 21)
        text_color_hex = self._design.resolve_color(
            style.get("text_color", "dark_gray")
        )
        body_font = style.get("font", self._design.get_font("body"))
        spacing_after = style.get("spacing_after", 80)

        self._render_list_items(
            doc, list_block.items,
            is_numbered=is_numbered,
            indent_per_level=indent_per_level,
            bullet_color_hex=bullet_color_hex,
            text_size_hp=text_size_hp,
            text_color_hex=text_color_hex,
            body_font=body_font,
            spacing_after=spacing_after,
            level=0,
            counter_start=1,
        )

        logger.debug(
            "Rendered %s list: %d item(s)",
            list_block.list_type.value, len(list_block.items),
        )

    def _render_list_items(
        self,
        doc: Document,
        items: list[ListItem],
        *,
        is_numbered: bool,
        indent_per_level: int,
        bullet_color_hex: str,
        text_size_hp: int,
        text_color_hex: str,
        body_font: str,
        spacing_after: int,
        level: int,
        counter_start: int,
    ) -> None:
        """Recursively render list items with proper nesting."""
        counter = counter_start

        for item in items:
            para = doc.add_paragraph()
            left_indent = indent_per_level * (level + 1)
            para.paragraph_format.left_indent = Twips(left_indent)
            para.paragraph_format.space_before = Twips(0)
            para.paragraph_format.space_after = Twips(spacing_after)

            # -- Bullet / number prefix --
            if is_numbered:
                prefix = f"{counter}. "
            else:
                prefix = "\u2022  "

            prefix_run = para.add_run(prefix)
            prefix_run.bold = True
            prefix_run.font.size = Pt(text_size_hp / 2.0)
            prefix_run.font.color.rgb = _hex_to_rgb(bullet_color_hex)
            prefix_run.font.name = body_font

            # -- Item text runs --
            if item.runs:
                for text_run in item.runs:
                    run = para.add_run(text_run.text)
                    run.font.size = Pt(text_size_hp / 2.0)
                    run.font.name = body_font
                    if text_run.bold:
                        run.bold = True
                    if text_run.italic:
                        run.italic = True
                    if text_run.color:
                        run.font.color.rgb = _hex_to_rgb(
                            self._design.resolve_color(text_run.color)
                        )
                    else:
                        run.font.color.rgb = _hex_to_rgb(text_color_hex)
            else:
                # Fallback to .text property if no runs are present
                text_content = item.text
                if text_content:
                    run = para.add_run(text_content)
                    run.font.size = Pt(text_size_hp / 2.0)
                    run.font.color.rgb = _hex_to_rgb(text_color_hex)
                    run.font.name = body_font

            # -- Nested children --
            if item.children:
                self._render_list_items(
                    doc, item.children,
                    is_numbered=is_numbered,
                    indent_per_level=indent_per_level,
                    bullet_color_hex=bullet_color_hex,
                    text_size_hp=text_size_hp,
                    text_color_hex=text_color_hex,
                    body_font=body_font,
                    spacing_after=spacing_after,
                    level=level + 1,
                    counter_start=1,
                )

            counter += 1

    # ── Steps rendering ───────────────────────────────────────────────

    def _render_steps(
        self,
        doc: Document,
        steps_block: StepsBlock,
        style: dict[str, Any],
    ) -> None:
        """Render a numbered procedure (steps block).

        Each step consists of:
        1. A large accent-coloured number.
        2. A bold title on the same line.
        3. A description paragraph in body style.
        """
        number_size_hp = style.get("number_size", 56)
        number_color_hex = self._design.resolve_color(
            style.get("number_color", "accent_blue")
        )
        number_bold = style.get("number_bold", True)

        title_size_hp = style.get("title_size", 24)
        title_bold = style.get("title_bold", True)
        title_color_hex = self._design.resolve_color(
            style.get("title_color", "black")
        )

        desc_size_hp = style.get("desc_size", 21)
        desc_color_hex = self._design.resolve_color(
            style.get("desc_color", "dark_gray")
        )

        display_font = style.get("display_font", self._design.get_font("display"))
        body_font = style.get("body_font", self._design.get_font("body"))
        spacing_between = style.get("spacing_between", 200)

        for step in steps_block.steps:
            # -- Step number + title paragraph --
            step_para = doc.add_paragraph()
            step_para.paragraph_format.space_before = Twips(spacing_between)
            step_para.paragraph_format.space_after = Twips(80)
            step_para.paragraph_format.keep_with_next = True

            # Number run
            num_run = step_para.add_run(f"{step.number:02d}")
            num_run.bold = number_bold
            num_run.font.size = Pt(number_size_hp / 2.0)
            num_run.font.color.rgb = _hex_to_rgb(number_color_hex)
            num_run.font.name = display_font

            # Separator
            sep_run = step_para.add_run("   ")
            sep_run.font.size = Pt(title_size_hp / 2.0)

            # Title run
            title_run = step_para.add_run(step.title)
            title_run.bold = title_bold
            title_run.font.size = Pt(title_size_hp / 2.0)
            title_run.font.color.rgb = _hex_to_rgb(title_color_hex)
            title_run.font.name = display_font

            # -- Description paragraph --
            if step.description or step.description_runs:
                desc_para = doc.add_paragraph()
                desc_para.paragraph_format.space_before = Twips(40)
                desc_para.paragraph_format.space_after = Twips(spacing_between)
                # Indent description to align with the title text
                desc_para.paragraph_format.left_indent = Twips(360)

                if step.description_runs:
                    for text_run in step.description_runs:
                        run = desc_para.add_run(text_run.text)
                        run.font.size = Pt(desc_size_hp / 2.0)
                        run.font.name = body_font
                        if text_run.bold:
                            run.bold = True
                        if text_run.italic:
                            run.italic = True
                        if text_run.color:
                            run.font.color.rgb = _hex_to_rgb(
                                self._design.resolve_color(text_run.color)
                            )
                        else:
                            run.font.color.rgb = _hex_to_rgb(desc_color_hex)
                else:
                    run = desc_para.add_run(step.description)
                    run.font.size = Pt(desc_size_hp / 2.0)
                    run.font.color.rgb = _hex_to_rgb(desc_color_hex)
                    run.font.name = body_font

        logger.debug(
            "Rendered steps block: %d step(s)", len(steps_block.steps)
        )

    # ── Page break rendering ──────────────────────────────────────────

    @staticmethod
    def _render_page_break(doc: Document) -> None:
        """Insert an explicit page break."""
        doc.add_page_break()
        logger.debug("Rendered page break")

    # ── Header / footer setup ─────────────────────────────────────────

    def _setup_header_footer(
        self,
        doc: Document,
        metadata: DocumentMetadata,
    ) -> None:
        """Configure running headers and footers on all document sections.

        Header: right-aligned, small text in light_gray showing the document
        title.

        Footer: left side shows "Confidentiel", right side shows the page
        number.  A thin top border separates the footer from body content.
        """
        header_cfg = self._design.get_header_config()
        footer_cfg = self._design.get_footer_config()

        for doc_section in doc.sections:
            # Allow different first-page header/footer so cover page
            # remains clean.
            doc_section.different_first_page_header_footer = True

            # ── Header ────────────────────────────────────────
            header = doc_section.header
            header.is_linked_to_previous = False

            header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            header_para.clear()

            alignment_str = header_cfg.get("alignment", "right")
            if alignment_str == "right":
                header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif alignment_str == "center":
                header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

            header_text = header_cfg.get("text", "") or metadata.title
            header_size_hp = header_cfg.get("size", 16)
            header_color_hex = self._design.resolve_color(
                header_cfg.get("color", "light_gray")
            )
            header_font = header_cfg.get("font", self._design.get_font("body"))

            if header_text:
                h_run = header_para.add_run(header_text)
                h_run.font.size = Pt(header_size_hp / 2.0)
                h_run.font.color.rgb = _hex_to_rgb(header_color_hex)
                h_run.font.name = header_font

            # ── Footer ────────────────────────────────────────
            footer = doc_section.footer
            footer.is_linked_to_previous = False

            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.clear()

            footer_size_hp = footer_cfg.get("size", 16)
            footer_color_hex = self._design.resolve_color(
                footer_cfg.get("color", "light_gray")
            )
            footer_font = footer_cfg.get("font", self._design.get_font("body"))

            border_top_size = footer_cfg.get("border_top_size", 1)
            border_top_color_hex = self._design.resolve_color(
                footer_cfg.get("border_top_color", "very_light_gray")
            )

            # Top border on the footer paragraph
            if border_top_size and border_top_size > 0:
                _add_top_border(footer_para, border_top_color_hex,
                                size_pt=border_top_size)

            # Use a tab stop to get left/right alignment within a single
            # paragraph.  Left text is "Confidentiel", right text is the
            # page number.
            left_text = footer_cfg.get("left_text", "Confidentiel")
            right_text_template = footer_cfg.get("right_text", "Page {page_number}")

            # Left portion
            left_run = footer_para.add_run(left_text)
            left_run.font.size = Pt(footer_size_hp / 2.0)
            left_run.font.color.rgb = _hex_to_rgb(footer_color_hex)
            left_run.font.name = footer_font

            # Tab to push the page number to the right
            tab_run = footer_para.add_run("\t")
            tab_run.font.size = Pt(footer_size_hp / 2.0)

            # Set a right-aligned tab stop at the usable width
            usable_width = self._design.usable_width
            pPr = footer_para._p.get_or_add_pPr()
            tabs_el = OxmlElement("w:tabs")
            tab_el = OxmlElement("w:tab")
            tab_el.set(qn("w:val"), "right")
            tab_el.set(qn("w:pos"), str(usable_width))
            tab_el.set(qn("w:leader"), "none")
            tabs_el.append(tab_el)
            pPr.append(tabs_el)

            # Right portion -- page number.
            # For a dynamic page number we use a PAGE field code.
            page_label = right_text_template.replace("{page_number}", "").strip()
            if page_label:
                page_label_run = footer_para.add_run(f"{page_label} ")
                page_label_run.font.size = Pt(footer_size_hp / 2.0)
                page_label_run.font.color.rgb = _hex_to_rgb(footer_color_hex)
                page_label_run.font.name = footer_font

            # Insert a PAGE field for the dynamic page number
            self._insert_page_number_field(
                footer_para, footer_size_hp, footer_color_hex, footer_font,
            )

        logger.debug("Header and footer configured for all sections")

    @staticmethod
    def _insert_page_number_field(
        para: DocxParagraph,
        size_hp: int,
        color_hex: str,
        font_name: str,
    ) -> None:
        """Insert a PAGE field code into *para* for dynamic page numbering.

        This creates the OOXML structure::

            <w:fldSimple w:instr=" PAGE ">
              <w:r><w:t>1</w:t></w:r>
            </w:fldSimple>
        """
        fld_simple = OxmlElement("w:fldSimple")
        fld_simple.set(qn("w:instr"), " PAGE ")

        # Create a run inside the field with the correct styling
        run_el = OxmlElement("w:r")

        # Run properties
        rPr = OxmlElement("w:rPr")
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(size_hp))
        rPr.append(sz)

        sz_cs = OxmlElement("w:szCs")
        sz_cs.set(qn("w:val"), str(size_hp))
        rPr.append(sz_cs)

        color_el = OxmlElement("w:color")
        color_el.set(qn("w:val"), color_hex.lstrip("#"))
        rPr.append(color_el)

        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), font_name)
        r_fonts.set(qn("w:hAnsi"), font_name)
        rPr.append(r_fonts)

        run_el.append(rPr)

        # Placeholder text (replaced by Word with the actual page number)
        text_el = OxmlElement("w:t")
        text_el.text = "1"
        run_el.append(text_el)

        fld_simple.append(run_el)
        para._p.append(fld_simple)
