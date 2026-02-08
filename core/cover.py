"""Cover page generator for the DocStyle Transformer.

Builds a styled cover page using python-docx, drawing layout parameters
from the design system (config/design-system.yaml) via the DesignSystem
mapper and document metadata from the intermediate representation.
"""

from __future__ import annotations

import logging
import re
from datetime import datetime
from typing import TYPE_CHECKING

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Inches, Pt, RGBColor, Twips

if TYPE_CHECKING:
    from docx.document import Document

from core.mapper import DesignSystem
from core.models import DocumentMetadata

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

_SPLIT_DELIMITERS = [" — ", " - ", " : "]


def _hex_to_rgb(hex_color: str) -> RGBColor:
    """Convert a ``#RRGGBB`` hex string to a python-docx *RGBColor*.

    The leading ``#`` is optional.
    """
    hex_color = hex_color.lstrip("#")
    if len(hex_color) != 6:
        raise ValueError(f"Invalid hex color: #{hex_color}")
    r = int(hex_color[0:2], 16)
    g = int(hex_color[2:4], 16)
    b = int(hex_color[4:6], 16)
    return RGBColor(r, g, b)


# ---------------------------------------------------------------------------
# CoverGenerator
# ---------------------------------------------------------------------------


class CoverGenerator:
    """Generates a formatted cover page and appends it to a *python-docx*
    ``Document``.

    Parameters
    ----------
    design : DesignSystem
        Provides colour resolution, cover-page configuration and font lookup
        sourced from the project design-system YAML.
    """

    def __init__(self, design: DesignSystem) -> None:
        self._design = design

    # -- public API ---------------------------------------------------------

    def generate(self, doc: Document, metadata: DocumentMetadata) -> None:
        """Add a complete cover page to *doc*.

        The cover page consists of (in order):

        1. A vertical spacer pushing the title block downward.
        2. A two-line title (line 1 black, line 2 accent blue).
        3. An accent bar (coloured bottom border).
        4. A subtitle line showing the version string.
        5. Bottom metadata (author, date, reference).
        6. A page break separating the cover from the body.
        """
        config = self._design.get_cover_config()
        logger.info("Generating cover page for '%s'", metadata.title)

        # Ensure a reference code exists.
        if not metadata.reference:
            metadata.reference = self._auto_reference(config)
            logger.debug("Auto-generated reference: %s", metadata.reference)

        self._add_top_spacer(doc, config)
        self._add_title(doc, metadata.title, config)
        self._add_accent_bar(doc, config)
        self._add_subtitle(doc, metadata.version, config)
        self._add_bottom_metadata(doc, metadata, config)
        self._add_page_break(doc)

        logger.info("Cover page generated successfully")

    # -- title splitting ----------------------------------------------------

    @staticmethod
    def _split_title(title: str) -> tuple[str, str]:
        """Intelligently split *title* into two roughly equal lines.

        Splitting strategy (in priority order):

        1. If the title contains an explicit delimiter (`` -- ``, `` - ``,
           `` : ``), split at the *first* delimiter closest to the midpoint.
        2. Otherwise split at the word boundary nearest to the midpoint.
        3. If the title is a single word, the second line is empty.
        """
        title = title.strip()
        if not title:
            return ("", "")

        mid = len(title) // 2

        # Strategy 1 -- explicit delimiters
        delimiter_positions: list[tuple[int, str]] = []
        for delim in _SPLIT_DELIMITERS:
            idx = title.find(delim)
            while idx != -1:
                delimiter_positions.append((idx, delim))
                idx = title.find(delim, idx + len(delim))

        if delimiter_positions:
            # Pick the delimiter whose position is closest to the midpoint.
            best_idx, best_delim = min(
                delimiter_positions, key=lambda t: abs(t[0] - mid)
            )
            line1 = title[:best_idx].rstrip()
            line2 = title[best_idx + len(best_delim):].lstrip()
            return (line1, line2)

        # Strategy 2 -- word boundary nearest to midpoint
        spaces = [i for i, ch in enumerate(title) if ch == " "]
        if not spaces:
            return (title, "")

        best_space = min(spaces, key=lambda idx: abs(idx - mid))
        return (title[:best_space], title[best_space + 1:])

    # -- private rendering helpers ------------------------------------------

    def _add_top_spacer(self, doc: Document, config: dict) -> None:
        """Insert empty paragraphs to push content down by *top_spacer* DXA."""
        spacer_dxa = config.get("top_spacer", 2400)
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Twips(spacer_dxa)
        para.paragraph_format.space_after = Twips(0)
        # Keep the paragraph visually empty.
        logger.debug("Top spacer: %d DXA", spacer_dxa)

    def _add_title(self, doc: Document, title: str, config: dict) -> None:
        """Render a two-line title block.

        The first line uses the primary title colour (black) and the second
        line uses the accent colour (accent_blue).
        """
        line1, line2 = self._split_title(title)
        title_size_hp = config.get("title_size", 144)  # half-points
        title_pt = title_size_hp / 2.0

        title_color_hex = self._design.resolve_color(
            config.get("title_color", "black")
        )
        accent_color_hex = self._design.resolve_color(
            config.get("title_accent_color", "accent_blue")
        )
        title_font = config.get("title_font") or self._design.get_font("display")

        # -- Line 1 (primary colour) --
        para1 = doc.add_paragraph()
        para1.paragraph_format.space_before = Twips(0)
        para1.paragraph_format.space_after = Twips(0)
        run1 = para1.add_run(line1)
        run1.font.size = Pt(title_pt)
        run1.font.color.rgb = _hex_to_rgb(title_color_hex)
        run1.font.bold = True
        if title_font:
            run1.font.name = title_font

        # -- Line 2 (accent colour) --
        if line2:
            para2 = doc.add_paragraph()
            para2.paragraph_format.space_before = Twips(0)
            para2.paragraph_format.space_after = Twips(0)
            run2 = para2.add_run(line2)
            run2.font.size = Pt(title_pt)
            run2.font.color.rgb = _hex_to_rgb(accent_color_hex)
            run2.font.bold = True
            if title_font:
                run2.font.name = title_font

        logger.debug(
            "Title rendered: line1=%r, line2=%r (%.1fpt)", line1, line2, title_pt
        )

    def _add_accent_bar(self, doc: Document, config: dict) -> None:
        """Insert a coloured accent bar implemented as a bottom border."""
        bar_color_hex = self._design.resolve_color(
            config.get("accent_bar_color", "accent_blue")
        )
        bar_width_dxa = config.get("accent_bar_width", 2400)
        bar_height_pt = config.get("accent_bar_height", 6)
        spacing_after = config.get("spacing_after", 200)

        # The bar is an otherwise empty paragraph with a bottom border.
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Twips(120)
        para.paragraph_format.space_after = Twips(spacing_after)

        # Apply an indent so the bar only spans *bar_width_dxa* from the left.
        usable_width = 9360  # default from design system
        right_indent_dxa = max(usable_width - bar_width_dxa, 0)
        para.paragraph_format.right_indent = Twips(right_indent_dxa)

        # Build the border XML element.
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        # Border size is in eighth-points.
        bottom.set(qn("w:sz"), str(bar_height_pt * 8))
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), bar_color_hex.lstrip("#"))
        pBdr.append(bottom)
        pPr.append(pBdr)

        logger.debug(
            "Accent bar: %d DXA wide, %dpt thick, color %s",
            bar_width_dxa,
            bar_height_pt,
            bar_color_hex,
        )

    def _add_subtitle(self, doc: Document, version: str, config: dict) -> None:
        """Render a subtitle line showing version information."""
        if not version:
            logger.debug("No version string provided; skipping subtitle")
            return

        subtitle_size_hp = config.get("subtitle_size", 32)
        subtitle_pt = subtitle_size_hp / 2.0
        subtitle_color_hex = self._design.resolve_color(
            config.get("subtitle_color", "medium_gray")
        )
        body_font = self._design.get_font("body")

        para = doc.add_paragraph()
        para.paragraph_format.space_before = Twips(200)
        para.paragraph_format.space_after = Twips(0)
        run = para.add_run(f"Version {version}")
        run.font.size = Pt(subtitle_pt)
        run.font.color.rgb = _hex_to_rgb(subtitle_color_hex)
        if body_font:
            run.font.name = body_font

        logger.debug("Subtitle: 'Version %s' at %.1fpt", version, subtitle_pt)

    def _add_bottom_metadata(
        self, doc: Document, metadata: DocumentMetadata, config: dict
    ) -> None:
        """Render the author, date and reference at the bottom of the cover."""
        meta_size_hp = config.get("metadata_size", 24)
        meta_pt = meta_size_hp / 2.0
        meta_color_hex = self._design.resolve_color(
            config.get("metadata_color", "light_gray")
        )
        meta_spacing = config.get("metadata_spacing", 160)
        body_font = self._design.get_font("body")

        # Insert a generous spacer to push metadata toward the bottom.
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Twips(3600)
        spacer.paragraph_format.space_after = Twips(0)

        lines: list[str] = []
        if metadata.author:
            lines.append(metadata.author)
        if metadata.date:
            lines.append(metadata.date)
        if metadata.reference:
            lines.append(metadata.reference)

        for line_text in lines:
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Twips(0)
            para.paragraph_format.space_after = Twips(meta_spacing)
            run = para.add_run(line_text)
            run.font.size = Pt(meta_pt)
            run.font.color.rgb = _hex_to_rgb(meta_color_hex)
            if body_font:
                run.font.name = body_font

        logger.debug("Bottom metadata: %d line(s)", len(lines))

    @staticmethod
    def _add_page_break(doc: Document) -> None:
        """Append a page break to separate the cover from document body."""
        doc.add_page_break()
        logger.debug("Page break added after cover")

    # -- reference generation -----------------------------------------------

    def _auto_reference(self, config: dict) -> str:
        """Generate a reference string using the configured format.

        The format string may contain ``{prefix}``, ``{year}`` and ``{seq}``
        placeholders.  *seq* defaults to ``1`` when auto-generating.
        """
        prefix = config.get("ref_prefix", "DOC")
        fmt = config.get("ref_format", "{prefix}-MAN-{year}-{seq:03d}")
        year = datetime.now().year
        try:
            ref = fmt.format(prefix=prefix, year=year, seq=1)
        except (KeyError, IndexError, ValueError) as exc:
            logger.warning(
                "Failed to format reference with template %r: %s", fmt, exc
            )
            ref = f"{prefix}-MAN-{year}-001"
        return ref
