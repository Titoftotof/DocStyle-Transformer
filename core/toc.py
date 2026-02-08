"""Table of contents page generator.

Generates a styled "Table des matières" (table of contents) page and inserts
it into a python-docx :class:`~docx.document.Document`.  The TOC layout uses
the project's :class:`DesignSystem` for consistent typography and colour.
"""

from __future__ import annotations

import logging
from typing import TYPE_CHECKING

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Twips

if TYPE_CHECKING:
    from docx.document import Document

from core.mapper import DesignSystem
from core.models import DocumentTree, Section

logger = logging.getLogger(__name__)

# ── Helpers ─────────────────────────────────────────────────────────────


def _hex_to_rgb(hex_color: str) -> RGBColor:
    """Convert a hex colour string (e.g. ``"#1F4E79"`` or ``"1F4E79"``) to an
    :class:`RGBColor` instance.

    Parameters
    ----------
    hex_color:
        A 6-digit hexadecimal colour with an optional leading ``#``.

    Returns
    -------
    RGBColor
        The corresponding python-docx colour value.
    """
    hex_color = hex_color.lstrip("#")
    r = int(hex_color[0:2], 16)
    g = int(hex_color[2:4], 16)
    b = int(hex_color[4:6], 16)
    return RGBColor(r, g, b)


# ── TOC generator ──────────────────────────────────────────────────────


class TocGenerator:
    """Builds a styled table of contents page inside a python-docx document.

    The generator walks the :class:`DocumentTree` section list and creates a
    TOC entry for every level-1 and level-2 section.  Level-1 entries are
    rendered in a prominent bold style while level-2 entries are indented and
    use a smaller, lighter treatment.

    Parameters
    ----------
    design:
        A :class:`DesignSystem` instance that provides colour resolution,
        font metadata, and component styles.

    Usage::

        toc = TocGenerator(design_system)
        toc.generate(doc, tree)
    """

    # Estimated starting page for TOC page numbering.  Because we do not have
    # access to the Word layout engine the page numbers are *approximations*.
    _ESTIMATED_TOC_START_PAGE = 3

    # Typographic constants
    _LEADER_CHAR = "."
    _LEADER_COUNT = 40  # default; adjusted per level

    def __init__(self, design: DesignSystem) -> None:
        self._design = design

    # ── Public API ────────────────────────────────────────────────────

    def generate(self, doc: Document, tree: DocumentTree) -> None:
        """Add a Table of Contents page to *doc* based on *tree*.

        The method appends:

        1. A "Table des matières" heading styled as Heading 1.
        2. An accent bar (bottom-bordered paragraph) under the heading.
        3. One paragraph per level-1 / level-2 section with number, title,
           dot-leader, and an estimated page number.
        4. A page break after the TOC.

        Parameters
        ----------
        doc:
            The python-docx :class:`Document` to append the TOC into.
        tree:
            The parsed :class:`DocumentTree` whose sections supply the TOC
            entries.
        """
        logger.info("Generating table of contents")

        sections = self._collect_toc_sections(tree)
        if not sections:
            logger.warning("No sections found for TOC generation; skipping")
            return

        # -- TOC heading --------------------------------------------------
        self._add_toc_heading(doc)

        # -- Accent bar under the heading ---------------------------------
        self._add_accent_bar(doc)

        # -- TOC entries --------------------------------------------------
        estimated_page = self._ESTIMATED_TOC_START_PAGE
        for section in sections:
            self._add_toc_entry(doc, section, estimated_page)
            # Rough page-count heuristic: each level-1 section ~2 pages,
            # level-2 sections don't bump the counter on their own.
            if section.level == 1:
                estimated_page += 2

        # -- Page break after TOC -----------------------------------------
        doc.add_page_break()
        logger.info(
            "TOC generated with %d entries",
            len(sections),
        )

    # ── Internal helpers ──────────────────────────────────────────────

    @staticmethod
    def _collect_toc_sections(tree: DocumentTree) -> list[Section]:
        """Return sections at level 1 and level 2 in document order."""
        return [s for s in tree.sections if s.level in (1, 2)]

    def _add_toc_heading(self, doc: Document) -> None:
        """Insert the "Table des matières" heading paragraph."""
        heading_style = self._design.get_heading_style(1)
        font_name = self._design.get_font("display")

        para = doc.add_paragraph()
        run = para.add_run("Table des matières")

        # Font styling
        run.bold = True
        run.font.size = Pt(heading_style.get("size", 40) / 2)
        run.font.name = font_name

        heading_color = heading_style.get("color")
        if heading_color:
            run.font.color.rgb = _hex_to_rgb(heading_color)

        # Paragraph spacing
        spacing_before = heading_style.get("spacing_before", 0)
        spacing_after = heading_style.get("spacing_after", 200)
        para.paragraph_format.space_before = Twips(spacing_before)
        para.paragraph_format.space_after = Twips(spacing_after)

        logger.debug("Added TOC heading paragraph")

    def _add_accent_bar(self, doc: Document) -> None:
        """Add a thin accent-coloured bottom border below the TOC heading.

        This creates an empty paragraph whose only purpose is to carry a
        bottom border via the underlying Open XML ``pBdr`` element.
        """
        accent_color = self._design.resolve_color("accent_blue")
        if accent_color:
            accent_hex = accent_color.lstrip("#")
        else:
            accent_hex = "1F4E79"

        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(16)

        # Build the border XML:
        #   <w:pBdr>
        #     <w:bottom w:val="single" w:sz="12" w:space="1" w:color="XXXXXX"/>
        #   </w:pBdr>
        p_pr = para._p.get_or_add_pPr()  # noqa: SLF001
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")  # border width in 1/8 pt
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), accent_hex)
        p_bdr.append(bottom)
        p_pr.append(p_bdr)

        logger.debug("Added accent bar under TOC heading")

    def _add_toc_entry(
        self,
        doc: Document,
        section: Section,
        estimated_page: int,
    ) -> None:
        """Add a single TOC entry paragraph for *section*.

        The entry is formatted as::

            NN — Section Title .............. PP

        where *NN* is the zero-padded section number and *PP* is the
        estimated page number.
        """
        # Resolve colours
        accent_color = self._design.resolve_color("accent_blue") or "#1F4E79"
        dark_gray = self._design.resolve_color("dark_gray") or "#404040"
        black = self._design.resolve_color("black") or "#000000"

        # Determine level-specific style
        is_level_1 = section.level == 1
        font_name = self._design.get_font("display" if is_level_1 else "body")
        font_size = Pt(13) if is_level_1 else Pt(11)
        title_color = _hex_to_rgb(black) if is_level_1 else _hex_to_rgb(dark_gray)

        # Section number (zero-padded to 2 digits)
        number_str = f"{section.number:02d}" if section.number is not None else "00"

        para = doc.add_paragraph()

        # -- Indentation for level 2 --------------------------------------
        if not is_level_1:
            para.paragraph_format.left_indent = Twips(720)  # ~0.5 inch

        para.paragraph_format.space_before = Pt(4 if is_level_1 else 2)
        para.paragraph_format.space_after = Pt(4 if is_level_1 else 2)

        # -- Section number run -------------------------------------------
        run_number = para.add_run(number_str)
        run_number.bold = True
        run_number.font.size = font_size
        run_number.font.name = font_name
        run_number.font.color.rgb = _hex_to_rgb(accent_color)

        # -- Separator run ------------------------------------------------
        run_sep = para.add_run(" \u2014 ")  # em-dash
        run_sep.bold = is_level_1
        run_sep.font.size = font_size
        run_sep.font.name = font_name
        run_sep.font.color.rgb = title_color

        # -- Title run ----------------------------------------------------
        run_title = para.add_run(section.heading)
        run_title.bold = is_level_1
        run_title.font.size = font_size
        run_title.font.name = font_name
        run_title.font.color.rgb = title_color

        # -- Dot leader + page number -------------------------------------
        leader_count = self._LEADER_COUNT - len(section.heading) // 2
        leader_count = max(leader_count, 6)
        leader = f"  {self._LEADER_CHAR * leader_count}  "

        run_leader = para.add_run(leader)
        run_leader.font.size = font_size
        run_leader.font.name = font_name
        run_leader.font.color.rgb = _hex_to_rgb(dark_gray)

        run_page = para.add_run(str(estimated_page))
        run_page.bold = is_level_1
        run_page.font.size = font_size
        run_page.font.name = font_name
        run_page.font.color.rgb = _hex_to_rgb(accent_color)

        logger.debug(
            "Added TOC entry: %s — %s (level %d, est. page %d)",
            number_str,
            section.heading,
            section.level,
            estimated_page,
        )
