"""Main application window for the DocStyle Transformer desktop UI.

Assembles the reusable widgets from :mod:`ui.components` into a complete
application.  The transformation is executed in a background thread so that
the UI remains responsive, with progress updates delivered through a
:class:`queue.Queue`.
"""

from __future__ import annotations

import logging
import queue
import threading
import tkinter as tk
from pathlib import Path
from tkinter import messagebox, ttk
from typing import Any, Optional

try:
    import ttkbootstrap as ttkb
    HAS_TTKB = True
except ImportError:
    HAS_TTKB = False

from core.detector import StructureDetector
from core.models import DocumentTree
from core.parser import DocxParser

from ui.components import (
    FileSelector,
    OptionsPanel,
    ProgressBar,
    ResultPanel,
    StructurePreview,
)

logger = logging.getLogger(__name__)

# ── Constants ────────────────────────────────────────────────────────

_APP_TITLE = "DocStyle Transformer"
_WINDOW_MIN_WIDTH = 800
_WINDOW_MIN_HEIGHT = 620
_POLL_INTERVAL_MS = 100  # how often the UI checks the progress queue


# ── Progress message types ────────────────────────────────────────────


class _ProgressMsg:
    """A message sent from the worker thread to the UI via the queue."""

    def __init__(
        self,
        kind: str,
        value: float = 0.0,
        label: str = "",
        data: Any = None,
    ) -> None:
        self.kind = kind      # "progress", "done", "error"
        self.value = value    # 0-100
        self.label = label    # human-readable status text
        self.data = data      # arbitrary payload (output path, stats, error msg)


# ── DocStyleApp ───────────────────────────────────────────────────────


class DocStyleApp:
    """Main application class that wires together all UI components.

    Usage::

        app = DocStyleApp()
        app.run()
    """

    def __init__(self) -> None:
        # ---- main window ----
        if HAS_TTKB:
            self.root = ttkb.Window(
                title=_APP_TITLE,
                themename="cosmo",
                minsize=(_WINDOW_MIN_WIDTH, _WINDOW_MIN_HEIGHT),
            )
        else:
            self.root = tk.Tk()
            self.root.title(_APP_TITLE)
            self.root.minsize(_WINDOW_MIN_WIDTH, _WINDOW_MIN_HEIGHT)

        # Centre on screen
        self.root.geometry(
            f"{_WINDOW_MIN_WIDTH}x{_WINDOW_MIN_HEIGHT}"
            f"+{(self.root.winfo_screenwidth() - _WINDOW_MIN_WIDTH) // 2}"
            f"+{(self.root.winfo_screenheight() - _WINDOW_MIN_HEIGHT) // 2}"
        )

        # ---- state ----
        self._parser = DocxParser()
        self._detector = StructureDetector()
        self._current_tree: Optional[DocumentTree] = None
        self._progress_queue: queue.Queue[_ProgressMsg] = queue.Queue()
        self._worker_thread: Optional[threading.Thread] = None

        # ---- build UI ----
        self._build_layout()

    # ── Layout ────────────────────────────────────────────────────────

    def _build_layout(self) -> None:
        """Create and arrange all widgets inside the main window."""
        pad = {"padx": 10, "pady": 5}

        # -- top: file selector --
        self._file_selector = FileSelector(
            self.root,
            label="Input .docx:",
            on_select=self._on_file_selected,
        )
        self._file_selector.pack(fill=tk.X, **pad)

        # -- middle: options + structure preview side by side --
        middle_frame = ttk.Frame(self.root)
        middle_frame.pack(fill=tk.BOTH, expand=True, **pad)

        self._options_panel = OptionsPanel(middle_frame)
        self._options_panel.pack(side=tk.LEFT, fill=tk.Y, padx=(0, 5))

        self._structure_preview = StructurePreview(middle_frame)
        self._structure_preview.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # -- transform button --
        btn_frame = ttk.Frame(self.root)
        btn_frame.pack(fill=tk.X, **pad)

        if HAS_TTKB:
            self._transform_btn = ttkb.Button(
                btn_frame,
                text="Transform",
                command=self._on_transform,
                bootstyle="success",
                state=tk.DISABLED,
            )
        else:
            self._transform_btn = ttk.Button(
                btn_frame,
                text="Transform",
                command=self._on_transform,
                state=tk.DISABLED,
            )
        self._transform_btn.pack(side=tk.LEFT)

        # -- progress bar --
        self._progress_bar = ProgressBar(btn_frame)
        self._progress_bar.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(10, 0))

        # -- bottom: result panel --
        self._result_panel = ResultPanel(self.root)
        self._result_panel.pack(fill=tk.X, **pad)

    # ── File selection callback ───────────────────────────────────────

    def _on_file_selected(self, path: str) -> None:
        """Parse the selected file and update the structure preview.

        Called by the :class:`FileSelector` widget when the user picks a
        new ``.docx`` file.
        """
        self._result_panel.reset()
        self._progress_bar.reset()
        self._current_tree = None
        self._transform_btn.configure(state=tk.DISABLED)

        if not path or not Path(path).is_file():
            return

        self._progress_bar.set_progress(10, "Parsing document...")

        try:
            tree = self._parser.parse(path)
            tree = self._detector.detect(tree)
        except Exception as exc:
            logger.exception("Failed to parse %s", path)
            messagebox.showerror(
                "Parse Error",
                f"Could not parse the selected file:\n\n{exc}",
                parent=self.root,
            )
            self._progress_bar.reset()
            return

        self._current_tree = tree
        self._structure_preview.update(tree)
        self._progress_bar.set_progress(100, "Document parsed successfully")
        self._transform_btn.configure(state=tk.NORMAL)

        logger.info(
            "Parsed '%s': %d sections, %d preamble elements",
            path,
            tree.section_count(),
            len(tree.preamble),
        )

    # ── Transform callback ────────────────────────────────────────────

    def _on_transform(self) -> None:
        """Kick off the transformation in a background thread.

        Disables the transform button, resets results, and starts polling
        the progress queue for updates from the worker.
        """
        input_path = self._file_selector.file_path
        if not input_path or self._current_tree is None:
            return

        # Prevent double-clicks
        if self._worker_thread is not None and self._worker_thread.is_alive():
            return

        self._transform_btn.configure(state=tk.DISABLED)
        self._result_panel.reset()
        self._progress_bar.set_progress(0, "Starting transformation...")

        options = self._options_panel.get_options()

        self._worker_thread = threading.Thread(
            target=self._transform_worker,
            args=(input_path, self._current_tree, options),
            daemon=True,
        )
        self._worker_thread.start()

        # Begin polling the queue
        self.root.after(_POLL_INTERVAL_MS, self._poll_progress)

    # ── Background worker ─────────────────────────────────────────────

    def _transform_worker(
        self,
        input_path: str,
        tree: DocumentTree,
        options: dict[str, Any],
    ) -> None:
        """Run the transformation pipeline in a background thread.

        Communicates progress back to the UI via ``self._progress_queue``.
        """
        q = self._progress_queue

        try:
            q.put(_ProgressMsg("progress", 10, "Loading design system..."))

            from core.mapper import DesignSystem, StyleMapper

            theme_path = options.get("theme_path") or None
            design = DesignSystem(theme_path=theme_path)
            mapper = StyleMapper(design)

            q.put(_ProgressMsg("progress", 30, "Mapping styles..."))

            # Map all elements to verify the style system works
            for section in tree.sections:
                mapper.map_heading(section.level)
                for child in section.children:
                    mapper.map_element(child)
            for elem in tree.preamble:
                mapper.map_element(elem)

            q.put(_ProgressMsg("progress", 50, "Generating output document..."))

            # Build output path
            input_p = Path(input_path)
            output_dir = input_p.parent
            output_name = input_p.stem + "_transformed.docx"
            output_path = str(output_dir / output_name)

            # Generate the output document
            from docx import Document as new_docx

            doc = new_docx()

            q.put(_ProgressMsg("progress", 60, "Building cover page..."))

            # Cover page
            if options.get("generate_cover", True):
                from core.cover import CoverGenerator

                cover_gen = CoverGenerator(design)
                metadata = tree.metadata
                if options.get("cover_title_override"):
                    metadata.title = options["cover_title_override"]
                cover_gen.generate(doc, metadata)

            q.put(_ProgressMsg("progress", 70, "Building table of contents..."))

            # TOC placeholder
            if options.get("generate_toc", True):
                toc_para = doc.add_paragraph("Table of Contents")
                toc_para.style = doc.styles["Heading 1"]
                for section in tree.sections:
                    indent = "    " * (section.level - 1)
                    label = section.heading
                    if options.get("number_sections") and section.number is not None:
                        label = f"{section.number:02d}  {label}"
                    doc.add_paragraph(f"{indent}{label}")
                # Page break after TOC
                toc_break = doc.add_paragraph()
                run = toc_break.add_run()
                run.add_break(docx_break_type=7)

            q.put(_ProgressMsg("progress", 85, "Writing sections..."))

            # Document body
            for section in tree.sections:
                heading_text = section.heading
                if options.get("number_sections") and section.number is not None:
                    heading_text = f"{section.number:02d}  {heading_text}"
                doc.add_heading(heading_text, level=min(section.level, 4))

                for child in section.children:
                    self._render_element(doc, child)

            # Mention
            if options.get("mention"):
                doc.add_paragraph("")
                doc.add_paragraph(options["mention"])

            q.put(_ProgressMsg("progress", 95, "Saving document..."))

            doc.save(output_path)

            stats = tree.summary()
            q.put(_ProgressMsg(
                "done", 100, "Transformation complete",
                data={"output_path": output_path, "stats": stats},
            ))

        except Exception as exc:
            logger.exception("Transformation failed")
            q.put(_ProgressMsg(
                "error", 0, "Transformation failed",
                data=str(exc),
            ))

    @staticmethod
    def _render_element(doc: Any, elem: Any) -> None:
        """Render a single content element into the python-docx document.

        This is a simplified renderer that produces basic output.  A
        production implementation would use the full style mapper and
        generator pipeline.
        """
        from core.models import (
            Callout,
            Image,
            ListBlock,
            PageBreak,
            Paragraph,
            StepsBlock,
            Table,
        )

        if isinstance(elem, Paragraph):
            text = elem.text
            if text.strip():
                doc.add_paragraph(text)

        elif isinstance(elem, Table):
            if elem.headers or elem.rows:
                cols = len(elem.headers) if elem.headers else (
                    len(elem.rows[0]) if elem.rows else 1
                )
                tbl = doc.add_table(
                    rows=1 + len(elem.rows), cols=cols
                )
                tbl.style = "Table Grid"
                # Headers
                if elem.headers:
                    for ci, hdr in enumerate(elem.headers):
                        tbl.rows[0].cells[ci].text = hdr
                # Data rows
                for ri, row in enumerate(elem.rows):
                    for ci, cell in enumerate(row):
                        if ci < cols:
                            tbl.rows[ri + 1].cells[ci].text = cell

        elif isinstance(elem, Callout):
            prefix = f"[{elem.callout_type.value.upper()}]"
            title = f" {elem.title}" if elem.title else ""
            doc.add_paragraph(f"{prefix}{title}: {elem.body}")

        elif isinstance(elem, ListBlock):
            for item in elem.items:
                bullet = "-" if elem.list_type.value == "bullet" else f"{elem.items.index(item) + 1}."
                indent = "  " * item.level
                doc.add_paragraph(f"{indent}{bullet} {item.text}")

        elif isinstance(elem, StepsBlock):
            for step in elem.steps:
                doc.add_paragraph(f"Step {step.number}: {step.title}")
                if step.description:
                    doc.add_paragraph(f"    {step.description}")

        elif isinstance(elem, Image):
            doc.add_paragraph(f"[Image: {elem.filename}]")

        elif isinstance(elem, PageBreak):
            para = doc.add_paragraph()
            run = para.add_run()
            run.add_break(docx_break_type=7)

    # ── Progress polling ──────────────────────────────────────────────

    def _poll_progress(self) -> None:
        """Drain the progress queue and update the UI accordingly.

        Reschedules itself until a terminal message (*done* or *error*)
        is received.
        """
        keep_polling = True

        while True:
            try:
                msg: _ProgressMsg = self._progress_queue.get_nowait()
            except queue.Empty:
                break

            if msg.kind == "progress":
                self._progress_bar.set_progress(msg.value, msg.label)

            elif msg.kind == "done":
                self._progress_bar.set_progress(100, msg.label)
                data = msg.data or {}
                self._result_panel.show_success(
                    output_path=data.get("output_path", ""),
                    stats=data.get("stats", {}),
                )
                self._transform_btn.configure(state=tk.NORMAL)
                keep_polling = False

            elif msg.kind == "error":
                self._progress_bar.set_progress(0, msg.label)
                self._result_panel.show_failure(str(msg.data))
                self._transform_btn.configure(state=tk.NORMAL)
                keep_polling = False

        if keep_polling:
            self.root.after(_POLL_INTERVAL_MS, self._poll_progress)

    # ── Public entry point ────────────────────────────────────────────

    def run(self) -> None:
        """Start the Tk main loop."""
        logger.info("Starting %s", _APP_TITLE)
        self.root.mainloop()


# ── Convenience launcher ──────────────────────────────────────────────


def main() -> None:
    """Entry point for running the application directly."""
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
    )
    app = DocStyleApp()
    app.run()


if __name__ == "__main__":
    main()
